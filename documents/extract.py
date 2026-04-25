"""
Text extraction for documents (Drive-hosted or local).

The assistant's read_document tool calls extract_text_from_drive() or
extract_text_from_local() and gets back a dict envelope:

    {
        "text": str,
        "truncated": bool,
        "warning": str | None,
        "total_chars": int,
        "offset": int,
        "next_offset": int | None,   # where to resume if truncated
    }
    {"error": str}

Callers pass an `offset` to resume reading past a prior truncation cut-off.
The full document is always re-extracted; the offset just controls which
slice is returned.

Supported formats:
- application/pdf                                                  (pypdf)
- application/vnd.openxmlformats-officedocument.wordprocessingml.document   (python-docx)
- application/vnd.openxmlformats-officedocument.spreadsheetml.sheet         (openpyxl)
- application/vnd.google-apps.document     -> exported to text/plain
- application/vnd.google-apps.spreadsheet  -> exported to text/csv
- application/vnd.google-apps.presentation -> exported to text/plain
- text/plain, text/csv, text/markdown                              (decoded utf-8)

Anything else returns {"error": "Unsupported mime type: ..."}.
"""

import logging
import os
from io import BytesIO

logger = logging.getLogger(__name__)

# Per-call slice cap on extracted text — about 50k tokens. Sized against
# Sonnet/Opus 200k context windows; models paginate via `offset` for docs
# larger than this.
MAX_CHARS = 200_000

# Per-sheet row cap for spreadsheets so a 50-sheet workbook doesn't blow up.
MAX_XLSX_ROWS_PER_SHEET = 1000

PDF_MIME = "application/pdf"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"

GOOGLE_MIME_EXPORTS = {
    "application/vnd.google-apps.document":     "text/plain",
    "application/vnd.google-apps.spreadsheet":  "text/csv",
    "application/vnd.google-apps.presentation": "text/plain",
}

PLAINTEXT_MIMES = {"text/plain", "text/csv", "text/markdown"}

# Map common file extensions to mime types for the local-file path.
EXT_TO_MIME = {
    ".pdf": PDF_MIME,
    ".docx": DOCX_MIME,
    ".xlsx": XLSX_MIME,
    ".txt": "text/plain",
    ".csv": "text/csv",
    ".md": "text/markdown",
}


def extract_text_from_drive(file_id, mime_type, offset=0):
    """Fetch a Drive file and extract its text. Returns dict envelope.

    `offset` lets callers resume past a prior truncation cut-off; the full
    document is re-downloaded and re-extracted either way.
    """
    from . import gdrive

    if not mime_type:
        return {"error": "Document is missing a mime type — cannot determine how to read it."}

    if mime_type in GOOGLE_MIME_EXPORTS:
        export_mime = GOOGLE_MIME_EXPORTS[mime_type]
        data = gdrive.export_file_bytes(file_id, export_mime)
        if data is None:
            return {"error": f"Failed to export Google file from Drive (mime: {mime_type})."}
        return _wrap_plaintext(_safe_decode(data), offset=offset)

    data = gdrive.download_file_bytes(file_id)
    if data is None:
        return {"error": "Failed to download file from Drive."}
    return _extract_by_mime(data, mime_type, offset=offset)


def extract_text_from_local(file_path, offset=0):
    """Read a local file from disk and extract its text. Returns dict envelope."""
    if not os.path.exists(file_path):
        return {"error": f"Local file not found at {file_path}"}
    ext = os.path.splitext(file_path)[1].lower()
    mime = EXT_TO_MIME.get(ext)
    if not mime:
        return {"error": f"Unsupported file extension: {ext}"}
    try:
        with open(file_path, "rb") as f:
            data = f.read()
    except Exception as e:
        logger.exception("Failed to read local file %s", file_path)
        return {"error": f"Failed to read local file: {e}"}
    return _extract_by_mime(data, mime, offset=offset)


def _extract_by_mime(data, mime_type, offset=0):
    """Dispatch raw bytes to the right extractor based on mime type."""
    if mime_type == PDF_MIME:
        return _wrap_plaintext(_extract_pdf(data), offset=offset)
    if mime_type == DOCX_MIME:
        return _wrap_plaintext(_extract_docx(data), offset=offset)
    if mime_type == XLSX_MIME:
        return _wrap_plaintext(_extract_xlsx(data), offset=offset)
    if mime_type in PLAINTEXT_MIMES:
        return _wrap_plaintext(_safe_decode(data), offset=offset)
    return {"error": f"Unsupported mime type: {mime_type}"}


def _wrap_plaintext(text, offset=0):
    """Apply offset/truncation and scanned/empty warnings to extracted text.

    Returns the slice text[offset:offset+MAX_CHARS]. When that slice doesn't
    reach the end, `truncated` is True and `next_offset` points at the first
    unread character so the caller can paginate.
    """
    if text is None:
        return {"error": "Text extraction returned no result."}

    total_chars = len(text)
    offset = max(0, int(offset))

    if total_chars and offset >= total_chars:
        return {
            "text": "",
            "truncated": False,
            "warning": (
                f"Offset {offset} is past the end of the document "
                f"(total: {total_chars} chars)."
            ),
            "total_chars": total_chars,
            "offset": offset,
            "next_offset": None,
        }

    end = offset + MAX_CHARS
    slice_text = text[offset:end]
    truncated = end < total_chars
    next_offset = end if truncated else None

    if truncated:
        slice_text += (
            f"\n\n[truncated at char {end} of {total_chars} — "
            f"call read_document with offset={next_offset} to continue reading]"
        )

    warning = None
    if offset == 0 and not slice_text.strip():
        warning = (
            "No extractable text — likely a scanned PDF, image-based document, "
            "or empty file. OCR is not available."
        )

    return {
        "text": slice_text,
        "truncated": truncated,
        "warning": warning,
        "total_chars": total_chars,
        "offset": offset,
        "next_offset": next_offset,
    }


def _safe_decode(data):
    """Decode bytes as UTF-8 with a permissive fallback."""
    if isinstance(data, str):
        return data
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("utf-8", errors="replace")


def _extract_pdf(data):
    """Extract concatenated text from all pages of a PDF."""
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.error("pypdf is not installed — cannot extract PDF text.")
        return ""

    try:
        reader = PdfReader(BytesIO(data))
    except Exception:
        logger.exception("Failed to open PDF")
        return ""

    pages = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception:
            logger.exception("Failed to extract text from PDF page %d", i)
            pages.append("")
    return "\n\n".join(p for p in pages if p)


def _extract_docx(data):
    """Extract paragraph and table text from a DOCX file."""
    try:
        import docx
    except ImportError:
        logger.error("python-docx is not installed — cannot extract DOCX text.")
        return ""

    try:
        document = docx.Document(BytesIO(data))
    except Exception:
        logger.exception("Failed to open DOCX")
        return ""

    parts = []
    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_xlsx(data):
    """Extract sheet contents from an XLSX file as CSV-ish text per sheet."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        logger.error("openpyxl is not installed — cannot extract XLSX text.")
        return ""

    wb = None
    try:
        wb = load_workbook(BytesIO(data), data_only=True, read_only=True)
    except Exception:
        logger.exception("Failed to open XLSX")
        return ""

    parts = []
    try:
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Sheet: {sheet_name} ===")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= MAX_XLSX_ROWS_PER_SHEET:
                    parts.append(f"...[truncated after {MAX_XLSX_ROWS_PER_SHEET} rows]")
                    break
                cells = ["" if c is None else str(c) for c in row]
                if any(cells):
                    parts.append(", ".join(cells))
            parts.append("")
    finally:
        try:
            wb.close()
        except Exception:
            pass
    return "\n".join(parts)
